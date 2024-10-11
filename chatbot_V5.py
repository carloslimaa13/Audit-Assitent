import streamlit as st
import PyPDF2
from pdfminer.high_level import extract_text
import re
from io import BytesIO
from docx import Document

# Título da página
st.write("""
# Audit Brasil
Bem-vindo ao assistente Audit Brasil.
Por aqui você pode ter um auxílio e otimizar seu tempo analisando contratos. Basta adicionar os contratos abaixo e digitar quais palavras ou frases deseja buscar dentro deles.
""")

# Função para limpar o texto
def limpar_texto(texto):
    texto = re.sub(r'\n\s*\n', '\n\n', texto)  # Mantém duas quebras de linha entre parágrafos
    texto = re.sub(r'-\s+', '', texto)  # Remove hífens no final de linha
    texto = re.sub(r'\b(?:Página\s*\d+|\d+\s*de\s*\d+)\b', ' ', texto)  # Remover cabeçalhos/rodapés genéricos
    texto = texto.replace('\u200b', '').strip()  # Remover caracteres invisíveis
    texto = texto.encode('utf-8', 'ignore').decode('utf-8')  # Remover caracteres especiais
    return texto

# Função para extrair texto do PDF
def extrair_texto_pdf(arquivo_carregado):
    try:
        reader = PyPDF2.PdfReader(arquivo_carregado)
        texto_extraido = ""
        for pagina in reader.pages:
            texto_extraido += pagina.extract_text()
        if texto_extraido:
            return limpar_texto(texto_extraido)
    except Exception as e:
        st.error(f"ERRO: Falha na extração com PyPDF2: {e}")
    try:
        texto_extraido = extract_text(arquivo_carregado)
        return limpar_texto(texto_extraido)
    except Exception as e:
        st.error(f"ERRO: Falha na extração com pdfminer: {e}")
    return None

# Função para buscar palavras ou frases nos textos extraídos
def buscar_texto(palavra, textos):
    resultados = {}
    padrao = rf'\b{re.escape(palavra)}\b'
    for nome_arquivo, texto in textos.items():
        paragrafos = texto.split('\n\n')
        encontrados = [paragrafo for paragrafo in paragrafos if re.search(padrao, paragrafo, re.IGNORECASE)]
        if encontrados:
            resultados[nome_arquivo] = encontrados
    return resultados

# Função para criar um arquivo Word com os parágrafos encontrados
def gerar_arquivo_word(paragrafos, nome_arquivo):
    doc = Document()
    doc.add_heading(f"Resultados para {nome_arquivo}", 0)
    for paragrafo in paragrafos:
        doc.add_paragraph(paragrafo)
    
    # Salvar o arquivo em um buffer de memória
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Volta ao início do buffer para o download
    return buffer

# Upload de múltiplos arquivos
st.write("### Adicione os arquivos abaixo")
arquivos_carregados = st.file_uploader("Arraste ou selecione os arquivos.", type="pdf", accept_multiple_files=True, help="Faça o upload aqui!")

# Dicionário para armazenar o conteúdo extraído de cada arquivo
textos_pdf = {}

# Criando a barra de progresso
progress_bar = st.progress(0)
total_arquivos = len(arquivos_carregados)

# Processar cada arquivo PDF carregado
if arquivos_carregados:
    for i, arquivo_carregado in enumerate(arquivos_carregados):
        try:
            texto_extraido = extrair_texto_pdf(arquivo_carregado)
            if texto_extraido:
                textos_pdf[arquivo_carregado.name] = texto_extraido
            else:
                st.error(f"ERRO: Erro ao extrair o texto de {arquivo_carregado.name}.")
        except Exception as e:
            st.error(f"ERRO: Falha ao processar {arquivo_carregado.name}: {e}")
        
        progresso_atual = (i + 1) / total_arquivos
        progress_bar.progress(progresso_atual)

    st.success("Todos os arquivos foram processados com sucesso!")

# Campo para o usuário digitar a palavra ou frase que deseja buscar
palavra_busca = st.text_input("Digite a palavra ou frase que deseja buscar:")

# Botão para realizar a busca
if st.button("Buscar"):
    if palavra_busca and textos_pdf:
        with st.spinner("Realizando a busca..."):
            resultados_busca = buscar_texto(palavra_busca, textos_pdf)
            if resultados_busca:
                for nome_arquivo, paragrafos in resultados_busca.items():
                    with st.expander(f"Resultados para '{palavra_busca}' em {nome_arquivo}"):
                        for paragrafo in paragrafos:
                            st.write(paragrafo)
                        
                        # Gerar arquivo Word com os parágrafos encontrados
                        buffer = gerar_arquivo_word(paragrafos, nome_arquivo)
                        
                        # Adiciona um botão de download para o arquivo Word
                        st.download_button(
                            label="Baixar resultados como Word",
                            data=buffer,
                            file_name=f"resultados_{nome_arquivo}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
            else:
                st.write("Nenhum resultado encontrado.")
    else:
        st.write("Por favor, insira uma palavra/frase válida e carregue arquivos PDF.")
